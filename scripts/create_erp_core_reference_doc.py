from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_PATH = "outputs/korcarvia_erp_core_reference.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True, color="0B2545")
        set_cell_shading(header_cells[idx], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
    if widths:
        set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run("\n" + body)
    for r in p.runs[1:]:
        r.font.name = "Malgun Gothic"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        r.font.size = Pt(9)
    doc.add_paragraph()


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("코카비아 ERP 구조 및 개발 기준서")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("공장 운영 흐름, 정산 기준, 확장 기능 구현 방향")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("기준일: 2026-06-22 / 용도: ERP 개발 기준 및 업무 설명서")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)

    add_callout(
        doc,
        "핵심 원칙",
        "ERP는 복잡한 프로그램이 아니라 공장의 실제 업무 기준을 그대로 기록하고 세는 도구다. "
        "개발 기준은 항상 업무 기준을 먼저 고정한 뒤, 그 기준이 코드와 데이터에서 어디에 반영되는지 확인하는 순서로 잡는다.",
    )

    doc.add_heading("1. 공장 업무의 큰 흐름", level=1)
    add_numbered(
        doc,
        [
            "차량 입고: 거래처 차량 또는 고객 직접 방문 차량이 들어온다.",
            "작업등록: 차량 전체 사진, 작업내용, 고객 요청사항을 입력한다.",
            "보험 선견적: ERP에 저장된 사진을 AOS 보험견적 프로그램에 올리고 보험사에 선견적을 보낸다.",
            "수리 협의 및 이슈 정리: 보험사와 수리 범위, 금액, 이슈를 협의하고 기록한다.",
            "수리 진행: 입고현황에서 수리 중인 차량을 관리한다.",
            "출고 예정 관리: 출고리스트에서 출고 예정일 기준으로 정리한다.",
            "출고 및 청구: 출고되면 출고일이 입력되고 보험사, 고객, 바디케어 센터 등에 청구한다.",
            "미결관리: 돈이 들어오기 전까지 미결, 미수, 장기미결, 장기미수를 관리한다.",
            "입금 확인 및 완결: 차량정산에서 수리비, 부가세, 면책금 등 입금 내역을 확인하고 완결 처리한다.",
            "추가 지출 및 종결: 입고지원비, 탁송비, 하자 수리비 등 사후 비용까지 입력한 뒤 종결 처리한다.",
        ],
    )

    doc.add_heading("2. ERP 화면별 역할", level=1)
    add_table(
        doc,
        ["화면", "업무 목적", "핵심 데이터"],
        [
            ["작업등록", "입고 차량을 최초 등록하고 사진, 작업내용, 고객 요청사항을 저장", "작업명, 차량번호, 사진, 요청사항, 보험/거래처 정보"],
            ["입고현황", "수리 중인 차량과 보험사 협의 이슈를 확인", "입고일, 작업상태, 담당자, 이슈 메모"],
            ["출고리스트", "출고 예정 및 출고 완료 차량을 날짜별로 관리", "출고예정일, 출고일, 작업명"],
            ["차량정산", "차량별 청구, 입금, 면책금, 수리비, 부가세를 기록", "청구일, 청구금액, 입금일, 입금금액, 계정"],
            ["미결관리", "출고 후 돈이 들어오기 전까지 관리할 차량을 분류", "진행상황, 출고일, 청구일, 입금일, 입금금액"],
            ["청구처별 리스트", "청구처별로 미결 차량과 금액을 묶어서 관리", "보험사/캐피탈/고객/바디케어, 작업명 기준 관리대수"],
            ["일일입출금", "매일 실제 들어오고 나가는 돈을 계정별로 기록", "일자, 계정, 입금, 지출, 잔고"],
            ["문서관리", "근태, 지출결의서, 경위서 등 내부 결재 문서 관리", "상태, 신청자, 부서, 승인자, 처리일"],
        ],
        [Cm(3.0), Cm(6.0), Cm(7.0)],
    )

    doc.add_heading("3. 정산과 미결관리 기준", level=1)
    add_callout(
        doc,
        "미결관리의 기준",
        "공통 기준은 출고일이 있고 진행상황이 미결인 차량이다. 관리대수는 작업명 기준 1건으로 센다. "
        "같은 작업명에 자차/대물/보험/캐피탈 등 여러 청구 내역이 있어도 관리대수는 1대다.",
    )
    add_table(
        doc,
        ["분류", "업무 기준", "해석"],
        [
            ["미결", "진행상황 미결이며 아직 입금으로 정리되지 않은 기본 관리 대상", "출고 후 청구/입금 확인이 필요한 차량"],
            ["미수", "입금일은 없고 입금금액이 있는 건", "실제 돈이 아직 확인되지 않았거나 미수 계정으로 잡힌 건"],
            ["장기미결", "청구일 기준 90일 초과", "오래된 청구 미수령 또는 지연 관리 대상"],
            ["장기미수", "미수 조건이면서 청구일 기준 90일 초과", "미수로 잡힌 상태가 장기화된 위험 건"],
            ["완결", "차량 관련 청구/입금 확인이 끝난 상태", "기본 정산 완료"],
            ["종결", "완결 이후 입고지원비, 탁송비, 하자 수리비 등 사후 지출까지 정리 완료", "최종 마감"],
        ],
        [Cm(3.0), Cm(6.0), Cm(7.0)],
    )

    doc.add_heading("4. 돈의 흐름과 계정 구조", level=1)
    add_table(
        doc,
        ["계정/통장", "사용 목적", "연동/주의 기준"],
        [
            ["국민은행", "중점적으로 사용하는 법인 통장", "보험사 입금, 일반 입금, 주요 잔고 기준"],
            ["부산은행", "BNK캐피탈 지급금 입금 계정", "캐피탈 관련 입금 구분 필요"],
            ["카드매출", "면책금, 수리비 등 카드 결제 발생", "입금일과 카드매출 처리일 구분 필요"],
            ["현대 BLUE POINT", "수리비 또는 면책금이 포인트로 발생", "BLUE POINT 계정으로 별도 집계"],
            ["법인1층 통장", "기존 보험사 입금 계정", "기존 데이터와 혼동되지 않게 계정명 유지"],
            ["현금", "현금 입출금 관리", "실제 보유 현금 잔고와 일일입출금 일치 필요"],
        ],
        [Cm(3.5), Cm(5.5), Cm(7.0)],
    )

    doc.add_heading("5. 일일입출금 연동 기준", level=1)
    add_bullets(
        doc,
        [
            "차량정산, 면책금관리, 지출결의서에서 돈이 발생해도 일일입출금에는 오늘 새로 입력된 값만 연동한다.",
            "입금일이나 지출일이 어제 또는 다른 날짜여도, 오늘 새로 입력한 값이면 오늘 일일입출금에 반영한다.",
            "과거에 입력된 기존 값을 오늘 다시 끌고 오면 계정별 잔고와 전체 잔고가 틀어진다.",
            "오늘 입력한 금액을 수정하는 경우 경고창을 띄운 뒤, 오늘 생성된 기존 연동값을 삭제하고 새 값으로 다시 반영한다.",
            "데이터 원본은 지우지 않고, 일일입출금의 잘못된 연동 표시만 제거하는 방식이 안전하다.",
        ],
    )

    doc.add_heading("6. 구현하고 싶은 기능과 가능성", level=1)
    add_table(
        doc,
        ["기능", "구현 가능성", "구현 방법"],
        [
            [
                "고객 개인정보동의서 카톡 발송 및 서명 회수",
                "가능",
                "카카오 알림톡/친구톡 또는 링크 기반 전자서명 페이지를 만든다. 고객에게 고유 링크를 보내고, 서명 이미지를 서버에 저장한 뒤 작업명과 연결한다.",
            ],
            [
                "현장 직원 휴대폰 사진 촬영 및 서버 자동 저장",
                "가능",
                "모바일 ERP에서 작업명을 선택하고 카메라 촬영을 실행한다. 사진은 NAS 서버의 작업명 폴더에 저장하고 DB에는 경로, 촬영자, 촬영시간을 기록한다.",
            ],
            [
                "사진 분석으로 수리 진행도 파악",
                "가능하나 단계적 구현 필요",
                "초기에는 사진 분류와 촬영 누락 체크부터 시작한다. 이후 AI로 파손부위, 작업 전/중/후 상태, 진행률 추정을 붙인다.",
            ],
            [
                "직원 간 메신저",
                "가능",
                "작업명별 채팅방과 부서별 공지방을 만든다. NAS 내부 DB에 메시지를 저장하고, 모바일 알림 또는 카톡 알림으로 확장한다.",
            ],
            [
                "국민은행 입출금 연동",
                "가능하나 은행 API/스크래핑 방식 검토 필요",
                "공식 기업뱅킹 API가 가능하면 API를 우선한다. 불가능하면 파일 업로드 또는 제한적 자동 수집 방식으로 시작한다.",
            ],
            [
                "세금계산서 연동",
                "가능",
                "홈택스/세무 프로그램 연동 가능 여부를 확인한다. 초기에는 엑셀/CSV 가져오기, 이후 API 연동으로 확장한다.",
            ],
            [
                "ERP 알림을 카톡으로 연동",
                "가능",
                "ERP 내부 알림을 먼저 정확히 만든 뒤, 중요 이벤트만 카카오 알림톡으로 보낸다. 예: 결재 요청, 장기미결 발생, 고객 서명 요청.",
            ],
        ],
        [Cm(3.5), Cm(3.0), Cm(9.5)],
    )

    doc.add_heading("7. 기능 구현 우선순위", level=1)
    add_numbered(
        doc,
        [
            "NAS ERP 안정화: 로그인, DB 스키마, 사진 저장, 핵심 화면 조회를 안정화한다.",
            "업무 기준 고정: 작업등록, 출고, 정산, 미결관리, 일일입출금의 기준을 문서 기준과 맞춘다.",
            "사진 서버화: 휴대폰/태블릿 촬영 사진을 작업명 폴더에 자동 저장한다.",
            "카톡/전자서명: 개인정보동의서 링크 발송과 서명 회수를 붙인다.",
            "알림 고도화: ERP 내부 알림 기준을 정리하고 카카오 알림톡으로 확장한다.",
            "은행/세금계산서 연동: 먼저 파일 가져오기 방식으로 안정화하고, 이후 공식 API를 검토한다.",
            "AI 사진 분석: 사진 저장 구조가 안정된 뒤 진행도 분석, 누락 사진 체크, 파손부위 분류를 단계적으로 붙인다.",
        ],
    )

    doc.add_heading("8. 개발할 때 반드시 지킬 기준", level=1)
    add_bullets(
        doc,
        [
            "업무 기준을 먼저 한 줄로 고정하고 코드를 본다.",
            "작업명 기준 1건인지, 청구 내역 기준 여러 건인지 항상 먼저 구분한다.",
            "출고일, 청구일, 입금일, 입력일은 서로 다른 의미이므로 섞지 않는다.",
            "돈과 잔고에 영향을 주는 기능은 원본 데이터를 삭제하지 않고 연동 데이터만 조정한다.",
            "경고창은 잘못된 접근, 기존 금액 수정, 잔고 영향이 있는 동작에서 적극적으로 사용한다.",
            "NAS 테스트용과 Vercel 운영용 폴더를 절대 섞지 않는다.",
        ],
    )

    doc.add_heading("9. 기준 문장", level=1)
    add_callout(
        doc,
        "앞으로의 개발 기준",
        "ERP는 현장 업무를 복잡하게 만드는 프로그램이 아니라, 공장이 이미 하고 있는 일을 정확히 기록하고 빠르게 확인하게 하는 프로그램이다. "
        "새 기능은 반드시 작업등록, 출고, 정산, 미결, 일일입출금, 문서/알림 흐름 중 어디에 들어가는지 먼저 정하고 구현한다.",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("코카비아 ERP 구조 및 개발 기준서")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
